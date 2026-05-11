from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

from app.core.config import settings
from app.models.report import Report


def _get_style(document: Document, *candidates: str):
    for candidate in candidates:
        try:
            return document.styles[candidate]
        except KeyError:
            continue
    return None


def _clear_document_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _clean_markdown(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"^#{1,6}\s*", "", cleaned)
    cleaned = cleaned.replace("**", "")
    cleaned = cleaned.replace("__", "")
    cleaned = cleaned.replace("`", "")
    cleaned = cleaned.replace("â€¢", "-")
    cleaned = cleaned.replace("•", "-")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _set_run_font(
    run,
    *,
    size: int = 12,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    color: str = "000000",
) -> None:
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _set_spacing(paragraph, *, before: int = 0, after: int = 0, line_spacing: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _set_paragraph_text(
    paragraph,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    size: int = 12,
    color: str = "000000",
) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic, underline=underline, color=color)


def _highlight(paragraph) -> None:
    for run in paragraph.runs:
        run.font.highlight_color = 7


def _add_spacer(document: Document, after_pt: int) -> None:
    paragraph = document.add_paragraph()
    _set_spacing(paragraph, after=after_pt)


def _apply_header(report: Report, document: Document) -> None:
    for section in document.sections:
        paragraphs = section.header.paragraphs
        if len(paragraphs) >= 5:
            _set_paragraph_text(
                paragraphs[0],
                "LAUDO TECNICO PERICIAL - JUSTICA DO TRABALHO",
                size=12,
                color="808080",
            )
            _set_paragraph_text(paragraphs[3], f"PROCESSO: {report.process.process_number}", size=10, color="808080")


def _add_centered(
    document: Document,
    text: str,
    *,
    style=None,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    size: int = 12,
    color: str = "000000",
    highlight: bool = False,
):
    paragraph = document.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_text(
        paragraph,
        text,
        bold=bold,
        italic=italic,
        underline=underline,
        size=size,
        color=color,
    )
    if highlight:
        _highlight(paragraph)
    _set_spacing(paragraph, after=0)
    return paragraph


def _add_cover(document: Document, report: Report) -> None:
    no_spacing = _get_style(document, "No Spacing", "SemEspaamento")

    _add_centered(document, report.process.court.upper(), style=no_spacing, bold=True, size=12)
    _add_centered(
        document,
        f"DA {report.process.labor_court.upper()} DO POLO REGIONAL DE {report.process.city.upper()} - {report.process.state.upper()}",
        style=no_spacing,
        bold=True,
        size=12,
        highlight=True,
    )

    _add_spacer(document, 92)
    _add_centered(document, "LAUDO PERICIAL", bold=True, size=26, color="1F5F6F")

    _add_spacer(document, 64)
    for line in [
        f"PROCESSO: {report.process.process_number}",
        f"RECLAMANTE: {report.process.claimant}",
        f"RECLAMADA: {report.process.defendant}",
    ]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(1.65)
        _set_paragraph_text(paragraph, line.upper(), size=12)
        _highlight(paragraph)
        _set_spacing(paragraph, after=4)

    _add_spacer(document, 78)
    _add_centered(document, report.process.expert_name, bold=True, italic=True, underline=True, size=12)
    _add_centered(document, "Engenheiro(a) de Seguranca do Trabalho", style=no_spacing, size=10)
    _add_centered(document, "Perito(a) Judicial", style=no_spacing, size=10)
    _add_centered(document, report.process.expert_registry, style=no_spacing, size=10)

    _add_spacer(document, 74)
    _add_centered(
        document,
        f"{report.process.city} - {report.process.state}, {report.created_at.strftime('%d/%m/%Y')}",
        size=12,
        highlight=True,
    )


def _format_section_title(index: int, title: str) -> str:
    normalized = _clean_markdown(title)
    normalized = re.sub(r"^\d+(\.\d+)?\s*", "", normalized).strip()
    if not normalized:
        normalized = f"SECAO {index}"
    return f"{index}.0 {normalized.upper()}"


def _add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_text(paragraph, _clean_markdown(text), bold=True, size=12)
    _set_spacing(paragraph, before=8, after=4)


def _add_body_paragraph(document: Document, text: str) -> None:
    content = _clean_markdown(text)
    if not content:
        return

    list_style = _get_style(document, "List Paragraph", "PargrafodaLista")
    body_style = _get_style(document, "Body Text", "Corpodetexto")
    default_style = _get_style(document, "Default", "Normal")

    is_list = content.startswith("-")
    paragraph = document.add_paragraph(style=(list_style if is_list else body_style) or default_style)
    if is_list:
        paragraph.paragraph_format.left_indent = Inches(0.45)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_paragraph_text(paragraph, content, size=12)
    _set_spacing(paragraph, after=0, line_spacing=1.15)


def _iter_clean_lines(content: str):
    for raw_line in content.splitlines():
        line = _clean_markdown(raw_line)
        if not line:
            continue
        pieces = re.split(
            r"(?=\b(?:Processo|Reclamante|Reclamada|Perito|CREA|Registro|Tribunal|Vara|Cidade|Data|Local|Observacoes):)",
            line,
        )
        for piece in pieces:
            cleaned = piece.strip(" -")
            if cleaned:
                yield cleaned


def _add_report_sections(document: Document, report: Report) -> None:
    for index, section in enumerate(sorted(report.sections, key=lambda item: item.section_order), start=1):
        _add_heading(document, _format_section_title(index, section.title))
        for raw_line in _iter_clean_lines(section.content):
            _add_body_paragraph(document, raw_line)


def _add_closing_block(document: Document, report: Report) -> None:
    document.add_paragraph()

    city_line = document.add_paragraph()
    city_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_text(city_line, f"{report.process.city} - {report.process.state}, {report.created_at.strftime('%d/%m/%Y')}.")

    document.add_paragraph()
    no_spacing = _get_style(document, "No Spacing", "SemEspaamento")
    _add_centered(document, report.process.expert_name, bold=True, italic=True, underline=True, size=12)
    _add_centered(document, "Engenheiro(a) de Seguranca do Trabalho", style=no_spacing, size=10)
    _add_centered(document, "Perito(a) Judicial", style=no_spacing, size=10)
    _add_centered(document, report.process.expert_registry, style=no_spacing, size=10)


def _add_attachments_section(document: Document) -> None:
    _add_heading(document, "ANEXOS")
    _add_heading(document, "ANEXO 01 - RELATORIO FOTOGRAFICO")
    _add_body_paragraph(document, "As fotos abaixo ilustram os locais periciados.")


def build_docx(report: Report) -> bytes:
    template_path = Path(settings.REPORT_TEMPLATE_PATH)
    document = Document(str(template_path)) if template_path.exists() else Document()

    _apply_header(report, document)
    _clear_document_body(document)
    _add_cover(document, report)
    document.add_page_break()
    _add_report_sections(document, report)
    _add_closing_block(document, report)
    document.add_page_break()
    _add_attachments_section(document)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf(report: Report) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 2 * cm

    def write_line(text: str, font: str = "Helvetica", size: int = 11):
        nonlocal y
        if y < 2 * cm:
            pdf.showPage()
            y = height - 2 * cm
        pdf.setFont(font, size)
        pdf.drawString(2 * cm, y, text[:100])
        y -= 0.6 * cm

    write_line("LaudoIA Pericial", "Helvetica-Bold", 16)
    write_line(report.title, "Helvetica-Bold", 13)
    write_line(f"Processo: {report.process.process_number}")
    write_line(f"Perito: {report.process.expert_name}")
    y -= 0.4 * cm

    for section in sorted(report.sections, key=lambda item: item.section_order):
        write_line(f"{section.section_order}. {section.title}", "Helvetica-Bold", 12)
        for line in section.content.splitlines() or [""]:
            write_line(line)
        y -= 0.2 * cm

    pdf.save()
    return buffer.getvalue()
